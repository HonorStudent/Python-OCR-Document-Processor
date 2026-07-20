"""
OCR Screenshots to DOCX

This Python program extracts text from image files using Optical Character
Recognition (OCR) and exports the extracted text to a Microsoft Word (.docx)
document.

Usage:
    1. Open a terminal or command prompt.
    2. Change to the directory containing this program.
    3. Run:

        python ocr_screenshots_to_docx.py "<path_to_images>"

Example:
    python ocr_screenshots_to_docx.py "C:\\Users\\YourName\\Pictures\\Screenshots"

The input folder can be any directory containing supported image files.
"""
#!/usr/bin/env python3
import argparse
from pathlib import Path

import pytesseract
from PIL import Image, ImageOps, ImageFilter
from docx import Document

# Optional imports
try:
    import cv2
    # Check if OpenCV has the needed constants
    if not hasattr(cv2, "ADAPTIVE_THRESH_GAUSSIAN_"):
        print("[INFO] Incomplete OpenCV build detected — using Pillow preprocessing instead.")
        cv2 = None
except Exception:
    cv2 = None

try:
    import pillow_heif  # for HEIC/HEIF support
    pillow_heif.register_heif_opener()
except Exception:
    pass

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp", ".heic", ".heif"}

def load_image(path: Path) -> Image.Image:
    return Image.open(path).convert("RGB")

def preprocess_pillow(img: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(img)
    gray = ImageOps.autocontrast(gray)
    gray = gray.filter(ImageFilter.UnsharpMask(radius=1, percent=80, threshold=3))
    return gray

def preprocess_opencv(img: Image.Image) -> Image.Image:
    import numpy as np
    arr = np.array(img)
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    gray = cv2.fastNlMeansDenoising(gray, None, 15, 7, 21)
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_,
                                cv2.THRESH_BINARY, 35, 11)
    return Image.fromarray(thr)

def ocr_image(img: Image.Image, lang: str, psm: int | None, oem: int | None) -> str:
    config_parts = []
    if psm is not None:
        config_parts.append(f"--psm {psm}")
    if oem is not None:
        config_parts.append(f"--oem {oem}")
    config = " ".join(config_parts)
    text = pytesseract.image_to_string(img, lang=lang, config=config)
    return text.strip()

def iter_images(folder: Path, recursive: bool):
    files = folder.rglob("*") if recursive else folder.glob("*")
    for f in files:
        if f.is_file() and f.suffix.lower() in IMAGE_EXTS:
            yield f

def main():
    ap = argparse.ArgumentParser(description="Extract text from screenshots/images and export to a Word document.")
    ap.add_argument("input_folder", help="Folder containing screenshots/images")
    ap.add_argument("-o", "--output", default="ocr_output.docx", help="Output .docx path (default: ocr_output.docx)")
    ap.add_argument("--lang", default="eng", help="Tesseract language(s), e.g. 'eng' or 'eng+tgl' (default: eng)")
    ap.add_argument("--recursive", action="store_true", help="Scan subfolders recursively")
    ap.add_argument("--per-file-page", dest="per_file_page", action="store_true",
                    help="Put each image's text on a new page")
    ap.add_argument("--include-filename-headings", dest="include_filename_headings", action="store_true",
                    help="Add a heading with the image filename before text")
    ap.add_argument("--save-txt", dest="save_txt", action="store_true",
                    help="Also save a .txt per image next to the image")
    ap.add_argument("--psm", type=int, default=None, help="Tesseract page segmentation mode, e.g., 6 or 3")
    ap.add_argument("--oem", type=int, default=None, help="Tesseract OCR Engine Mode (0..3)")
    args = ap.parse_args()

    input_dir = Path(args.input_folder)
    if not input_dir.is_dir():
        raise SystemExit(f"Not a directory: {input_dir}")

    doc = Document()
    doc.add_heading("OCR Export", level=0)
    doc.add_paragraph(f"Source folder: {input_dir.resolve()}")
    doc.add_paragraph(f"Language(s): {args.lang}")

    count = 0
    for img_path in sorted(iter_images(input_dir, args.recursive), key=lambda p: p.name.lower()):
        try:
            img = load_image(img_path)
            pre = preprocess_opencv(img) if cv2 is not None else preprocess_pillow(img)
            text = ocr_image(pre, lang=args.lang, psm=args.psm, oem=args.oem)
        except Exception as e:
            text = f"[OCR ERROR: {e}]"

        if args.per_file_page and count > 0:
            doc.add_page_break()

        if args.include_filename_headings:
            doc.add_heading(img_path.name, level=2)

        para = doc.add_paragraph()
        para.add_run(text if text else "[No text detected]")

        if args.save_txt and text:
            try:
                txt_out = img_path.with_suffix(img_path.suffix + ".txt")
                txt_out.write_text(text, encoding="utf-8")
            except Exception:
                pass

        count += 1

    doc.save(args.output)
    print(f"Processed {count} image(s). Wrote: {Path(args.output).resolve()}")

if __name__ == "__main__":
    main()
